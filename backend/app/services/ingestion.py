import html
import logging
import re
from collections.abc import Callable
from concurrent.futures import Future, ThreadPoolExecutor
from dataclasses import dataclass, field
from io import BytesIO

import fitz

from app.config import Settings, get_settings
from app.services.file_types import detect_file_type
from app.services.ocr import OCRService
from app.services.pdf_embedded_images import (
    EmbeddedFigure,
    ParsedAttachedAsset,
    _figure_to_attached_asset,
    attach_figures_to_chunks,
    extract_figures_from_page,
    figure_bboxes_on_page,
    figure_overlaps_bboxes,
)
from app.services.pdf_captions import apply_page_caption_matches, extract_page_captions
from app.services.pdf_layout_regions import layout_region
from app.services.pdf_refs import attach_by_explicit_refs
from app.services.pdf_raster_tables import promote_figures_to_raster_tables
from app.services.pdf_table_merge import merge_cross_page_tables
from app.services.pdf_header_footer import HeaderFooterFilter, build_header_footer_filter
from app.services.pdf_tables import (
    EmbeddedTable,
    _table_to_attached_asset,
    attach_tables_to_chunks,
    extract_tables_from_page,
    filter_page_blocks,
    filter_page_text,
    table_bboxes_on_page,
)

logger = logging.getLogger(__name__)

_SECTION_PATH_MAX_LEN = 512
_SECTION_PATH_SEPARATOR = " > "
_PAGE_SEPARATOR = "\n\n"
_TOC_LEADER_RE = re.compile(r"\.{4,}|…{2,}|·{4,}")
_TOC_ENTRY_RE = re.compile(
    r"^(?:\d+(?:\.\d+)*)?\s*.+?"
    r"(?:\.{3,}|…{2,}|·{4,}|\s{2,})"
    r"\s*\d+\s*$"
)
_CHAPTER_ONE_RE = re.compile(
    r"第[一1壹]章|chapter\s*1\b",
    re.IGNORECASE,
)
_CHAPTER_TITLE_RE = re.compile(
    r"第[一二三四五六七八九十百千零两\d]+章|chapter\s+\d+",
    re.IGNORECASE,
)
_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")


BlockSpan = tuple[int, int, float, float]
PageRow = tuple[
    str | None,
    str,
    int,
    tuple[float, float, float, float] | None,
    list[BlockSpan] | None,
]


LayoutRegion = dict[str, int | list[float]]


@dataclass
class ParsedChunk:
    text: str
    page: int | None
    section: str | None
    chunk_index: int
    layout_bbox: tuple[float, float, float, float] | None = None
    layout_regions: list[LayoutRegion] | None = None
    sort_key: float | None = None
    layout_y1: float | None = None
    attached_assets: list[ParsedAttachedAsset] = field(default_factory=list)


@dataclass(frozen=True)
class TocEntry:
    level: int
    title: str
    start_page: int
    end_page: int
    path: str


@dataclass(frozen=True)
class TocAnchor:
    level: int
    title: str
    path: str
    page: int
    y: float
    has_y: bool = False


@dataclass
class ParseResult:
    chunks: list[ParsedChunk]
    page_count: int
    ocr_pages: int
    toc_entries: list[TocEntry]


def toc_entry_to_dict(entry: TocEntry) -> dict:
    return {
        "level": entry.level,
        "title": entry.title,
        "start_page": entry.start_page,
        "end_page": entry.end_page,
        "path": entry.path,
    }


def toc_entries_from_dicts(items: list[dict]) -> list[TocEntry]:
    return [
        TocEntry(
            level=int(item["level"]),
            title=str(item["title"]),
            start_page=int(item["start_page"]),
            end_page=int(item["end_page"]),
            path=str(item["path"]),
        )
        for item in items
    ]


def is_toc_text(text: str) -> bool:
    """Detect table-of-contents style text (section title + dot leaders + page number)."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if len(lines) < 2:
        return False

    matched = 0
    for line in lines:
        if _TOC_ENTRY_RE.match(line):
            matched += 1
            continue
        if _TOC_LEADER_RE.search(line) and re.search(r"\d+\s*$", line):
            matched += 1
            continue
        if re.fullmatch(r"[\.\s·…]{4,}\s*\d+", line):
            matched += 1

    return matched >= 2 and matched / len(lines) >= 0.4


def _split_text_with_offsets(
    text: str, chunk_size: int, overlap: int
) -> list[tuple[int, str]]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [(0, text)]

    chunks: list[tuple[int, str]] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            split_at = text.rfind("\n", start + chunk_size // 2, end)
            if split_at == -1:
                split_at = text.rfind("。", start + chunk_size // 2, end)
            if split_at == -1:
                split_at = text.rfind(". ", start + chunk_size // 2, end)
            if split_at != -1:
                end = split_at + 1
        piece = text[start:end].strip()
        if piece:
            chunks.append((start, piece))
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _block_spans_from_joined_blocks(
    texts: list[str],
    block_bounds: list[tuple[float, float]],
) -> list[BlockSpan]:
    spans: list[BlockSpan] = []
    offset = 0
    for index, (text, (y0, y1)) in enumerate(zip(texts, block_bounds, strict=True)):
        if index > 0:
            offset += 1
        start = offset
        offset += len(text)
        spans.append((start, offset, y0, y1))
    return spans


def _y_bounds_for_range(
    block_spans: list[BlockSpan],
    start: int,
    end: int,
) -> tuple[float, float] | None:
    y0_values: list[float] = []
    y1_values: list[float] = []
    for span_start, span_end, y0, y1 in block_spans:
        if span_end <= start or span_start >= end:
            continue
        y0_values.append(y0)
        y1_values.append(y1)
    if not y0_values:
        return None
    return (min(y0_values), max(y1_values))


def _layout_y_bounds(
    layout_bbox: tuple[float, float, float, float] | None,
) -> tuple[float | None, float | None]:
    if layout_bbox is None:
        return None, None
    return float(layout_bbox[1]), float(layout_bbox[3])


def _merge_bboxes(
    bboxes: list[tuple[float, float, float, float]],
) -> tuple[float, float, float, float] | None:
    if not bboxes:
        return None
    return (
        min(bbox[0] for bbox in bboxes),
        min(bbox[1] for bbox in bboxes),
        max(bbox[2] for bbox in bboxes),
        max(bbox[3] for bbox in bboxes),
    )


def _page_content_bbox(
    page: fitz.Page,
    exclude_bboxes: list[tuple[float, float, float, float]],
    hf: HeaderFooterFilter | None = None,
) -> tuple[float, float, float, float] | None:
    blocks = filter_page_blocks(page, exclude_bboxes, hf=hf)
    merged = _merge_bboxes([block[:4] for block in blocks])
    if merged:
        return merged
    rect = page.rect
    return (float(rect.x0), float(rect.y0), float(rect.x1), float(rect.y1))


def _group_contiguous_sections(
    pages: list[PageRow],
) -> list[tuple[str | None, list[tuple[int, str, tuple[float, float, float, float] | None, list[BlockSpan] | None]]]]:
    if not pages:
        return []

    groups: list[
        tuple[str | None, list[tuple[int, str, tuple[float, float, float, float] | None, list[BlockSpan] | None]]]
    ] = []
    current_section = pages[0][0]
    current_pages: list[tuple[int, str, tuple[float, float, float, float] | None, list[BlockSpan] | None]] = [
        (pages[0][2], pages[0][1], pages[0][3], pages[0][4])
    ]

    for section, text, page, bbox, block_spans in pages[1:]:
        if section != current_section:
            groups.append((current_section, current_pages))
            current_section = section
            current_pages = []
        current_pages.append((page, text, bbox, block_spans))
    groups.append((current_section, current_pages))
    return groups


def _concat_pages(
    pages: list[tuple[int, str, tuple[float, float, float, float] | None, list[BlockSpan] | None]],
    separator: str = _PAGE_SEPARATOR,
) -> tuple[str, list[tuple[int, int, tuple[float, float, float, float] | None]], list[BlockSpan]]:
    spans: list[tuple[int, int, tuple[float, float, float, float] | None]] = []
    block_spans: list[BlockSpan] = []
    parts: list[str] = []
    offset = 0
    for index, (page, text, bbox, page_block_spans) in enumerate(pages):
        if index > 0:
            offset += len(separator)
        spans.append((page, offset, bbox))
        if page_block_spans:
            for span_start, span_end, y0, y1 in page_block_spans:
                block_spans.append((offset + span_start, offset + span_end, y0, y1))
        parts.append(text)
        offset += len(text)
    return separator.join(parts), spans, block_spans


def _page_for_offset(spans: list[tuple[int, int, tuple[float, float, float, float] | None]], offset: int) -> int:
    page = spans[0][0]
    for span_page, start, _bbox in spans:
        if start <= offset:
            page = span_page
        else:
            break
    return page


def _bbox_for_offset(
    spans: list[tuple[int, int, tuple[float, float, float, float] | None]],
    offset: int,
) -> tuple[float, float, float, float] | None:
    bbox: tuple[float, float, float, float] | None = None
    for _page, start, span_bbox in spans:
        if start <= offset:
            bbox = span_bbox
        else:
            break
    return bbox


def _regions_for_range(
    page_spans: list[tuple[int, int, tuple[float, float, float, float] | None]],
    block_spans: list[BlockSpan],
    start: int,
    end: int,
    total_len: int,
    separator_len: int = len(_PAGE_SEPARATOR),
) -> list[LayoutRegion]:
    """Map a chunk character range to per-page layout bboxes (supports cross-page chunks)."""
    if not page_spans or start >= end:
        return []

    regions: list[LayoutRegion] = []
    for index, (page, page_start, page_bbox) in enumerate(page_spans):
        if page_bbox is None:
            continue
        if index + 1 < len(page_spans):
            page_end = page_spans[index + 1][1] - separator_len
        else:
            page_end = total_len

        overlap_start = max(start, page_start)
        overlap_end = min(end, page_end)
        if overlap_start >= overlap_end:
            continue

        y0_values: list[float] = []
        y1_values: list[float] = []
        for span_start, span_end, y0, y1 in block_spans:
            if span_end <= overlap_start or span_start >= overlap_end:
                continue
            if span_end <= page_start or span_start >= page_end:
                continue
            y0_values.append(y0)
            y1_values.append(y1)

        if y0_values:
            region_bbox = (page_bbox[0], min(y0_values), page_bbox[2], max(y1_values))
        else:
            region_bbox = page_bbox

        regions.append(layout_region(page, region_bbox))

    return regions


def _extract_native_page_text(
    page: fitz.Page,
    hf: HeaderFooterFilter | None = None,
) -> str:
    blocks = filter_page_blocks(page, [], hf=hf)
    if blocks:
        return "\n".join(text for *_, text in blocks)
    return ""


def _extract_page_text_blocks(page: fitz.Page) -> list[tuple[float, float, str]]:
    blocks = page.get_text("blocks")
    extracted: list[tuple[float, float, str]] = []
    for block in blocks:
        if len(block) < 5:
            continue
        bbox = fitz.Rect(block[:4])
        text = str(block[4]).strip()
        if text:
            extracted.append((float(bbox.y0), float(bbox.y1), text))
    extracted.sort(key=lambda item: (item[0], item[1]))
    return extracted


def _segment_page_by_anchors(
    anchors: list[TocAnchor],
    page_number: int,
    blocks: list[tuple[float, float, float, float, str]],
    fallback_section: str | None,
) -> list[tuple[str | None, str, tuple[float, float, float, float] | None, list[BlockSpan]]]:
    if not blocks:
        return []

    segments: list[
        tuple[str | None, list[str], list[tuple[float, float, float, float]], list[tuple[float, float]]]
    ] = []
    current_section: str | None = None
    current_texts: list[str] = []
    current_bboxes: list[tuple[float, float, float, float]] = []
    current_bounds: list[tuple[float, float]] = []

    for x0, y0, x1, y1, text in blocks:
        mid_y = (y0 + y1) / 2
        section = section_at_position(anchors, page_number, mid_y) or fallback_section
        block_bbox = (x0, y0, x1, y1)
        if section != current_section and current_texts:
            segments.append((current_section, current_texts, current_bboxes, current_bounds))
            current_texts = []
            current_bboxes = []
            current_bounds = []
        current_section = section
        current_texts.append(text)
        current_bboxes.append(block_bbox)
        current_bounds.append((y0, y1))

    if current_texts:
        segments.append((current_section, current_texts, current_bboxes, current_bounds))

    return [
        (
            section,
            "\n".join(texts),
            _merge_bboxes(bboxes),
            _block_spans_from_joined_blocks(texts, bounds),
        )
        for section, texts, bboxes, bounds in segments
        if texts
    ]


def _page_row_from_blocks(
    section: str | None,
    page_number: int,
    blocks: list[tuple[float, float, float, float, str]],
    layout_bbox: tuple[float, float, float, float] | None,
) -> PageRow | None:
    texts = [text for *_, text in blocks]
    if not texts:
        return None
    bounds = [(y0, y1) for _, y0, _, y1, _ in blocks]
    return (
        section,
        "\n".join(texts),
        page_number,
        layout_bbox,
        _block_spans_from_joined_blocks(texts, bounds),
    )


def _page_needs_ocr(native_text: str, settings: Settings) -> bool:
    if not settings.ocr_enabled:
        return False
    if settings.ocr_force:
        return True
    return len(native_text.strip()) < settings.ocr_min_chars_per_page


def _truncate_section_path(path: str) -> str:
    if len(path) <= _SECTION_PATH_MAX_LEN:
        return path
    return path[: _SECTION_PATH_MAX_LEN - 1] + "…"


def _dest_raw_point(dest: dict) -> tuple[float, float] | None:
    to = dest.get("to")
    if to is None:
        return None
    try:
        if hasattr(to, "x"):
            return float(to.x), float(to.y)
        return float(to[0]), float(to[1])
    except (TypeError, ValueError, IndexError):
        return None


def _dest_y_from_pdf_space(page: fitz.Page, raw_x: float, raw_y: float) -> float:
    pdf_point = fitz.Point(raw_x, raw_y)
    return float((pdf_point * page.transformation_matrix).y)


def _count_non_decreasing(values: list[float]) -> int:
    if len(values) < 2:
        return 0
    return sum(1 for index in range(len(values) - 1) if values[index] <= values[index + 1])


def _toc_dest_coordinate_votes(page: fitz.Page, raw_ys: list[float]) -> tuple[int, int]:
    """Score top-down vs PDF-space interpretation for one page's bookmark Y values.

    Returns (top_down_votes, pdf_space_votes). Multi-bookmark pages vote via
    monotonic pairs in outline order; single-bookmark pages add one zone vote.
    """
    if not raw_ys:
        return 0, 0

    page_height = float(page.rect.height)
    flipped_ys = [_dest_y_from_pdf_space(page, 0.0, raw_y) for raw_y in raw_ys]
    top_down_votes = _count_non_decreasing(raw_ys)
    pdf_votes = _count_non_decreasing(flipped_ys)

    if len(raw_ys) == 1:
        raw_y = raw_ys[0]
        flipped_y = flipped_ys[0]
        near_top_raw = raw_y <= page_height * 0.45
        near_bottom_flipped = flipped_y >= page_height * 0.55
        near_top_flipped = flipped_y <= page_height * 0.45
        near_bottom_raw = raw_y >= page_height * 0.55
        if near_top_raw and near_bottom_flipped:
            top_down_votes += 1
        elif near_top_flipped and near_bottom_raw:
            pdf_votes += 1

    return top_down_votes, pdf_votes


def _document_toc_dest_already_top_down(
    doc: fitz.Document,
    page_raw_ys: dict[int, list[float]],
) -> bool:
    """Vote across every page that has bookmark Y values in the outline."""
    top_down_total = 0
    pdf_total = 0
    for page_num, raw_ys in page_raw_ys.items():
        top_down_votes, pdf_votes = _toc_dest_coordinate_votes(doc[page_num - 1], raw_ys)
        top_down_total += top_down_votes
        pdf_total += pdf_votes
    if top_down_total != pdf_total:
        return top_down_total > pdf_total
    # Tie: keep PDF-space conversion for backward compatibility.
    return False


def _dest_to_page_y(
    page: fitz.Page,
    dest: dict,
    *,
    already_top_down: bool = False,
) -> float | None:
    """Convert a PDF bookmark destination to MuPDF top-left Y."""
    point = _dest_raw_point(dest)
    if point is None:
        return None
    raw_x, raw_y = point
    if already_top_down:
        return raw_y
    return _dest_y_from_pdf_space(page, raw_x, raw_y)


def _parse_raw_toc(doc: fitz.Document) -> list[tuple[int, str, int, float | None]]:
    raw_toc = doc.get_toc(simple=False)
    if not raw_toc:
        return []

    page_raw_ys: dict[int, list[float]] = {}
    pending_rows: list[tuple[int, str, int, dict | None, float | None]] = []

    for row in raw_toc:
        level = int(row[0])
        title = str(row[1]).strip()
        page = max(1, int(row[2]))
        dest = row[3] if len(row) > 3 else None
        raw_y: float | None = None
        if isinstance(dest, dict):
            dest_page = dest.get("page")
            if dest_page is not None and int(dest_page) >= 0:
                page = int(dest_page) + 1
            if 1 <= page <= doc.page_count:
                point = _dest_raw_point(dest)
                if point is not None:
                    raw_y = point[1]
                    page_raw_ys.setdefault(page, []).append(raw_y)
        pending_rows.append((level, title, page, dest, raw_y))

    document_already_top_down = _document_toc_dest_already_top_down(doc, page_raw_ys)

    parsed: list[tuple[int, str, int, float | None]] = []
    for level, title, page, dest, raw_y in pending_rows:
        y: float | None = None
        if isinstance(dest, dict) and raw_y is not None and 1 <= page <= doc.page_count:
            y = _dest_to_page_y(
                doc[page - 1],
                dest,
                already_top_down=document_already_top_down,
            )
        parsed.append((level, title, page, y))
    return parsed


def _build_toc_paths(
    parsed: list[tuple[int, str, int, float | None]],
) -> list[tuple[int, str, int, float | None, str]]:
    path_titles: list[str] = []
    path_levels: list[int] = []
    with_paths: list[tuple[int, str, int, float | None, str]] = []
    for level, title, page, y in parsed:
        while path_levels and path_levels[-1] >= level:
            path_levels.pop()
            path_titles.pop()
        path_titles.append(title)
        path_levels.append(level)
        path = _truncate_section_path(_SECTION_PATH_SEPARATOR.join(path_titles))
        with_paths.append((level, title, page, y, path))
    return with_paths


def _toc_anchors_from_parsed(
    parsed: list[tuple[int, str, int, float | None]],
) -> list[TocAnchor]:
    if not parsed:
        return []

    anchors: list[TocAnchor] = []
    for level, title, page, y, path in _build_toc_paths(parsed):
        anchors.append(
            TocAnchor(
                level=level,
                title=title,
                path=path,
                page=page,
                y=y if y is not None else 0.0,
                has_y=y is not None,
            )
        )
    return sorted(anchors, key=lambda anchor: (anchor.page, anchor.y, anchor.level))


def _toc_entries_from_parsed(
    parsed: list[tuple[int, str, int, float | None]],
    page_count: int,
) -> list[TocEntry]:
    if not parsed:
        return []

    normalized: list[tuple[int, str, int, int]] = []
    for index, (level, title, page, _y) in enumerate(parsed):
        start_page = max(1, int(page))
        end_page = page_count
        for next_index in range(index + 1, len(parsed)):
            next_level, _, next_page, _ = parsed[next_index]
            if next_level <= level:
                end_page = max(start_page, int(next_page) - 1)
                break
        normalized.append((int(level), title, start_page, end_page))

    path_titles: list[str] = []
    path_levels: list[int] = []
    entries: list[TocEntry] = []
    for level, title, start_page, end_page in normalized:
        while path_levels and path_levels[-1] >= level:
            path_levels.pop()
            path_titles.pop()
        path_titles.append(title)
        path_levels.append(level)
        path = _SECTION_PATH_SEPARATOR.join(path_titles)
        entries.append(
            TocEntry(
                level=level,
                title=title,
                start_page=start_page,
                end_page=end_page,
                path=_truncate_section_path(path),
            )
        )
    return entries


def parse_pdf_toc(doc: fitz.Document) -> tuple[list[TocEntry], list[TocAnchor]]:
    parsed = _parse_raw_toc(doc)
    if not parsed:
        return [], []
    return (
        _toc_entries_from_parsed(parsed, doc.page_count),
        _toc_anchors_from_parsed(parsed),
    )


def build_toc_anchors(doc: fitz.Document) -> list[TocAnchor]:
    _, anchors = parse_pdf_toc(doc)
    return anchors


def build_toc_entries(doc: fitz.Document) -> list[TocEntry]:
    entries, _ = parse_pdf_toc(doc)
    return entries


def section_at_position(anchors: list[TocAnchor], page: int, y: float) -> str | None:
    path_titles: list[str] = []
    path_levels: list[int] = []
    for anchor in anchors:
        if (anchor.page, anchor.y) > (page, y):
            break
        while path_levels and path_levels[-1] >= anchor.level:
            path_levels.pop()
            path_titles.pop()
        path_titles.append(anchor.title)
        path_levels.append(anchor.level)
    if not path_titles:
        return None
    return _truncate_section_path(_SECTION_PATH_SEPARATOR.join(path_titles))


def section_for_page(entries: list[TocEntry], page: int) -> str | None:
    matches = [entry for entry in entries if entry.start_page <= page <= entry.end_page]
    if not matches:
        return None
    return max(matches, key=lambda entry: entry.level).path


def first_chapter_start_page(entries: list[TocEntry]) -> int | None:
    """Return the PDF page where the first chapter begins, based on bookmarks."""
    if not entries:
        return None

    for entry in entries:
        if _CHAPTER_ONE_RE.search(entry.title) or _CHAPTER_ONE_RE.search(entry.path):
            return entry.start_page

    level_one = [entry for entry in entries if entry.level == 1]
    for entry in level_one:
        if _CHAPTER_TITLE_RE.search(entry.title):
            return entry.start_page

    if level_one:
        return min(entry.start_page for entry in level_one)

    return min(entry.start_page for entry in entries)


def should_skip_front_matter(page: int, toc_entries: list[TocEntry]) -> bool:
    content_start = first_chapter_start_page(toc_entries)
    return content_start is not None and page < content_start


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Unable to decode text file")


def _strip_html(raw_html: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "", raw_html)
    text = re.sub(r"(?s)<[^>]+>", "\n", text)
    text = html.unescape(text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _pages_with_bbox(
    pages: list[tuple[str | None, str, int]],
) -> list[PageRow]:
    return [(section, text, page, None, None) for section, text, page in pages]


def _flush_text_buffer(
    pages: list[tuple[str | None, str, int]],
    section: str | None,
    buffer: list[str],
    page_number: int = 1,
) -> None:
    if not buffer:
        return
    text = "\n".join(buffer).strip()
    if text and not is_toc_text(text):
        pages.append((section, text, page_number))


def _parse_structured_text_pages(text: str, *, markdown: bool) -> list[tuple[str | None, str, int]]:
    text = text.strip()
    if not text:
        return []

    if not markdown:
        return [(None, text, 1)]

    pages: list[tuple[str | None, str, int]] = []
    section: str | None = None
    buffer: list[str] = []
    for line in text.splitlines():
        match = _MD_HEADING_RE.match(line)
        if match:
            _flush_text_buffer(pages, section, buffer)
            section = match.group(2).strip()
            buffer = []
            continue
        buffer.append(line)
    _flush_text_buffer(pages, section, buffer)
    return pages if pages else [(None, text, 1)]


def _parse_docx_pages(file_bytes: bytes) -> list[tuple[str | None, str, int]]:
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(file_bytes))
    pages: list[tuple[str | None, str, int]] = []
    section: str | None = None
    buffer: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            _flush_text_buffer(pages, section, buffer)
            section = text
            buffer = []
            continue
        buffer.append(text)
    _flush_text_buffer(pages, section, buffer)

    if pages:
        return pages

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            pages.append((section, text, 1))
    return pages


def _merge_pdf_layout_chunks(
    inline_chunks: list[tuple[int, float, ParsedChunk]],
    page_chunks: list[ParsedChunk],
) -> list[ParsedChunk]:
    ordered: list[tuple[int, float, int, ParsedChunk]] = []

    for page, sort_key, chunk in inline_chunks:
        ordered.append((page, sort_key, 0, chunk))

    for chunk in page_chunks:
        ordered.append((chunk.page or 0, float("inf"), chunk.chunk_index + 1, chunk))

    ordered.sort(key=lambda item: (item[0], item[1], item[2]))
    merged: list[ParsedChunk] = []
    for index, (_, _, _, chunk) in enumerate(ordered):
        chunk.chunk_index = index
        merged.append(chunk)
    return merged


def _orphan_figure_chunk(figure: EmbeddedFigure) -> ParsedChunk:
    """Standalone figure chunk; caption stays on attached asset, not chunk.text."""
    return ParsedChunk(
        text="",
        page=figure.page,
        section=figure.section,
        chunk_index=0,
        layout_bbox=figure.bbox,
        layout_regions=[layout_region(figure.page, figure.bbox)],
        sort_key=float(figure.bbox[1]),
        layout_y1=float(figure.bbox[3]),
        attached_assets=[_figure_to_attached_asset(figure)],
    )


def _orphan_table_chunk(table: EmbeddedTable) -> ParsedChunk:
    """Standalone table chunk; summary stays on attached asset, not chunk.text."""
    regions = (
        list(table.layout_regions)
        if table.layout_regions
        else [layout_region(table.page, table.bbox)]
    )
    return ParsedChunk(
        text="",
        page=table.page,
        section=table.section,
        chunk_index=0,
        layout_bbox=table.bbox,
        layout_regions=regions,
        sort_key=float(table.bbox[1]),
        layout_y1=float(table.bbox[3]),
        attached_assets=[_table_to_attached_asset(table)],
    )


def _build_chunks_from_pages(
    pages: list[PageRow],
    settings: Settings,
) -> list[ParsedChunk]:
    parsed_chunks: list[ParsedChunk] = []
    chunk_index = 0
    for section, page_group in _group_contiguous_sections(pages):
        section_text, page_spans, block_spans = _concat_pages(page_group)
        for offset, piece in _split_text_with_offsets(
            section_text,
            settings.rag_chunk_size,
            settings.rag_chunk_overlap,
        ):
            chunk_end = offset + len(piece)
            layout_regions = _regions_for_range(
                page_spans,
                block_spans,
                offset,
                chunk_end,
                len(section_text),
            )
            layout_bbox = _bbox_for_offset(page_spans, offset)
            if layout_regions:
                first_bbox = layout_regions[0]["bbox"]
                if isinstance(first_bbox, list) and len(first_bbox) == 4:
                    layout_bbox = tuple(float(value) for value in first_bbox)
            y_bounds = _y_bounds_for_range(block_spans, offset, chunk_end)
            if y_bounds is not None:
                sort_key, layout_y1 = y_bounds
            else:
                sort_key, layout_y1 = _layout_y_bounds(layout_bbox)
            parsed_chunks.append(
                ParsedChunk(
                    text=piece,
                    page=_page_for_offset(page_spans, offset),
                    section=section,
                    chunk_index=chunk_index,
                    layout_bbox=layout_bbox,
                    layout_regions=layout_regions or None,
                    sort_key=sort_key,
                    layout_y1=layout_y1,
                )
            )
            chunk_index += 1
    return parsed_chunks


class IngestionService:
    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or get_settings()
        self.ocr = OCRService(self.settings) if self.settings.ocr_enabled else None

    def _render_page_for_ocr(self, page: fitz.Page) -> tuple[bytes, float]:
        scale = self.settings.ocr_render_scale
        matrix = fitz.Matrix(scale, scale)
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)
        return pixmap.tobytes("png"), float(page.rect.height) * scale

    def _recognize_rendered_page(
        self,
        png_bytes: bytes,
        page_height: float,
        hf: HeaderFooterFilter | None,
    ) -> str:
        if self.ocr is None:
            return ""
        return self.ocr.recognize_bytes(
            png_bytes,
            page_height=page_height,
            hf=hf,
        )

    def _extract_page_text_with_ocr_future(
        self,
        page: fitz.Page,
        *,
        hf: HeaderFooterFilter | None,
        exclude_bboxes: list[tuple[float, float, float, float]],
        ocr_future: Future[str] | None,
    ) -> tuple[str, bool]:
        native_text = _extract_native_page_text(page, hf=hf)

        if exclude_bboxes and not self.settings.ocr_force:
            return filter_page_text(page, exclude_bboxes, hf=hf), False

        used_ocr = False
        if _page_needs_ocr(native_text, self.settings) and self.ocr is not None:
            if ocr_future is not None:
                ocr_text = ocr_future.result()
            else:
                png_bytes, page_height = self._render_page_for_ocr(page)
                ocr_text = self._recognize_rendered_page(png_bytes, page_height, hf)
            if len(ocr_text.strip()) > len(native_text.strip()):
                native_text = ocr_text
                used_ocr = True

        page_text = native_text
        if exclude_bboxes and page_text.strip():
            page_text = filter_page_text(page, exclude_bboxes, hf=hf) or page_text
        return page_text, used_ocr

    def _append_page_text_rows(
        self,
        pages: list[PageRow],
        *,
        page: fitz.Page,
        page_number: int,
        page_text: str,
        used_ocr: bool,
        exclude_bboxes: list[tuple[float, float, float, float]],
        toc_anchors: list[TocAnchor],
        toc_entries: list[TocEntry],
        use_y_split: bool,
        hf_filter: HeaderFooterFilter | None,
    ) -> None:
        if not page_text.strip() and not exclude_bboxes:
            return
        if page_text.strip() and is_toc_text(page_text):
            return

        fallback_section = section_for_page(toc_entries, page_number)
        if use_y_split and not used_ocr and page_text.strip():
            blocks = filter_page_blocks(page, exclude_bboxes, hf=hf_filter)
            if blocks:
                for section, segment_text, segment_bbox, block_spans in _segment_page_by_anchors(
                    toc_anchors,
                    page_number,
                    blocks,
                    fallback_section,
                ):
                    if segment_text.strip():
                        pages.append(
                            (
                                section,
                                segment_text.strip(),
                                page_number,
                                segment_bbox,
                                block_spans,
                            )
                        )
            elif page_text.strip():
                pages.append(
                    (
                        fallback_section,
                        page_text.strip(),
                        page_number,
                        _page_content_bbox(page, exclude_bboxes, hf=hf_filter),
                        None,
                    )
                )
        elif page_text.strip():
            blocks = filter_page_blocks(page, exclude_bboxes, hf=hf_filter)
            row = _page_row_from_blocks(
                fallback_section,
                page_number,
                blocks,
                _page_content_bbox(page, exclude_bboxes, hf=hf_filter),
            )
            if row is not None:
                pages.append(row)

    def parse_pdf(
        self,
        file_bytes: bytes,
        on_page_progress: Callable[[int, int], None] | None = None,
    ) -> ParseResult:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = doc.page_count
        toc_entries, toc_anchors = parse_pdf_toc(doc)
        use_y_split = any(anchor.has_y for anchor in toc_anchors)
        skip_front_matter = lambda page_number: should_skip_front_matter(
            page_number, toc_entries
        )
        hf_filter = build_header_footer_filter(
            doc,
            self.settings,
            should_skip_page=skip_front_matter,
        )

        def section_resolver(page_number: int, y: float | None = None) -> str | None:
            if y is not None and use_y_split:
                resolved = section_at_position(toc_anchors, page_number, y)
                if resolved:
                    return resolved
            return section_for_page(toc_entries, page_number)

        pages: list[PageRow] = []
        ocr_pages = 0
        extracted_tables: list[EmbeddedTable] = []
        embedded_figures: list[EmbeddedFigure] = []
        png_cache: dict[int, tuple[bytes, int, int]] = {}
        seen_placements: set[tuple[int, int, tuple[int, int, int, int]]] = set()
        tables_available = True
        promote_processed = 0
        max_workers = max(1, self.settings.pdf_parallel_workers)

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            for page_index in range(page_count):
                page_number = page_index + 1
                if should_skip_front_matter(page_number, toc_entries):
                    continue

                page = doc[page_index]

                page_tables: list[EmbeddedTable] = []
                if tables_available:
                    try:
                        page_tables = extract_tables_from_page(
                            page,
                            page_number,
                            settings=self.settings,
                            section_resolver=section_resolver,
                        )
                    except AttributeError:
                        logger.warning(
                            "PyMuPDF find_tables is unavailable; table extraction disabled"
                        )
                        tables_available = False

                page_figures = extract_figures_from_page(
                    page,
                    page_number,
                    doc,
                    settings=self.settings,
                    section_resolver=section_resolver,
                    png_cache=png_cache,
                    seen_placements=seen_placements,
                )
                table_bboxes = table_bboxes_on_page(page_tables, page_number)
                page_figures = [
                    figure
                    for figure in page_figures
                    if not figure_overlaps_bboxes(figure, table_bboxes)
                ]

                page_figures, promoted, promote_processed = promote_figures_to_raster_tables(
                    page_figures,
                    settings=self.settings,
                    processed=promote_processed,
                )
                page_tables = [*page_tables, *promoted]

                page_captions = extract_page_captions(page, page_number)
                page_figures, page_tables = apply_page_caption_matches(
                    page_captions,
                    page_figures,
                    page_tables,
                )
                embedded_figures.extend(page_figures)
                extracted_tables.extend(page_tables)

                exclude_bboxes = (
                    table_bboxes_on_page(page_tables, page_number)
                    + figure_bboxes_on_page(page_figures, page_number)
                )

                native_text = _extract_native_page_text(page, hf=hf_filter)
                will_need_ocr = (
                    not (exclude_bboxes and not self.settings.ocr_force)
                    and _page_needs_ocr(native_text, self.settings)
                    and self.ocr is not None
                )
                ocr_future: Future[str] | None = None
                if will_need_ocr:
                    png_bytes, page_height = self._render_page_for_ocr(page)
                    ocr_future = executor.submit(
                        self._recognize_rendered_page,
                        png_bytes,
                        page_height,
                        hf_filter,
                    )

                page_text, used_ocr = self._extract_page_text_with_ocr_future(
                    page,
                    hf=hf_filter,
                    exclude_bboxes=exclude_bboxes,
                    ocr_future=ocr_future,
                )
                if used_ocr:
                    ocr_pages += 1

                self._append_page_text_rows(
                    pages,
                    page=page,
                    page_number=page_number,
                    page_text=page_text,
                    used_ocr=used_ocr,
                    exclude_bboxes=exclude_bboxes,
                    toc_anchors=toc_anchors,
                    toc_entries=toc_entries,
                    use_y_split=use_y_split,
                    hf_filter=hf_filter,
                )
                if on_page_progress is not None:
                    on_page_progress(page_number, page_count)

        if extracted_tables:
            page_heights = {
                page_index + 1: float(doc[page_index].rect.height)
                for page_index in range(page_count)
            }
            extracted_tables = merge_cross_page_tables(
                extracted_tables,
                page_heights=page_heights,
                settings=self.settings,
            )

        if not pages and not embedded_figures and not extracted_tables:
            doc.close()
            raise ValueError("No text extracted from PDF")

        page_chunks = _build_chunks_from_pages(pages, self.settings) if pages else []

        remaining_figures, remaining_tables = attach_by_explicit_refs(
            page_chunks,
            embedded_figures,
            extracted_tables,
        )

        orphan_tables = attach_tables_to_chunks(remaining_tables, page_chunks)

        orphan_figures = attach_figures_to_chunks(
            remaining_figures,
            page_chunks,
        )

        inline_chunks: list[tuple[int, float, ParsedChunk]] = []
        for table in orphan_tables:
            inline_chunks.append(
                (table.page, table.sort_key, _orphan_table_chunk(table))
            )
        for figure in orphan_figures:
            inline_chunks.append(
                (figure.page, figure.sort_key, _orphan_figure_chunk(figure))
            )

        parsed_chunks = (
            _merge_pdf_layout_chunks(inline_chunks, page_chunks)
            if inline_chunks
            else page_chunks
        )

        doc.close()
        return ParseResult(
            chunks=parsed_chunks,
            page_count=page_count,
            ocr_pages=ocr_pages,
            toc_entries=toc_entries,
        )

    def _parse_text_document(
        self,
        file_bytes: bytes,
        *,
        markdown: bool,
        on_page_progress: Callable[[int, int], None] | None = None,
    ) -> ParseResult:
        pages = _parse_structured_text_pages(_decode_text_bytes(file_bytes), markdown=markdown)
        if not pages:
            raise ValueError("No text extracted from file")
        if on_page_progress is not None:
            on_page_progress(1, 1)
        return ParseResult(
            chunks=_build_chunks_from_pages(_pages_with_bbox(pages), self.settings),
            page_count=1,
            ocr_pages=0,
            toc_entries=[],
        )

    def _parse_html_document(
        self,
        file_bytes: bytes,
        on_page_progress: Callable[[int, int], None] | None = None,
    ) -> ParseResult:
        text = _strip_html(_decode_text_bytes(file_bytes))
        if not text:
            raise ValueError("No text extracted from HTML")
        pages = [(None, text, 1)]
        if on_page_progress is not None:
            on_page_progress(1, 1)
        return ParseResult(
            chunks=_build_chunks_from_pages(_pages_with_bbox(pages), self.settings),
            page_count=1,
            ocr_pages=0,
            toc_entries=[],
        )

    def _parse_docx_document(
        self,
        file_bytes: bytes,
        on_page_progress: Callable[[int, int], None] | None = None,
    ) -> ParseResult:
        pages = _parse_docx_pages(file_bytes)
        if not pages:
            raise ValueError("No text extracted from Word document")
        if on_page_progress is not None:
            on_page_progress(1, 1)
        return ParseResult(
            chunks=_build_chunks_from_pages(_pages_with_bbox(pages), self.settings),
            page_count=1,
            ocr_pages=0,
            toc_entries=[],
        )

    def _parse_image_document(
        self,
        file_bytes: bytes,
        on_page_progress: Callable[[int, int], None] | None = None,
    ) -> ParseResult:
        if self.ocr is None:
            raise ValueError("OCR is disabled; cannot process image files")
        text = self.ocr.recognize_bytes(file_bytes).strip()
        if not text:
            raise ValueError("No text extracted from image")
        pages = [(None, text, 1)]
        if on_page_progress is not None:
            on_page_progress(1, 1)
        return ParseResult(
            chunks=_build_chunks_from_pages(pages, self.settings),
            page_count=1,
            ocr_pages=1,
            toc_entries=[],
        )

    def parse_document(
        self,
        file_bytes: bytes,
        filename: str,
        on_page_progress: Callable[[int, int], None] | None = None,
    ) -> ParseResult:
        file_type = detect_file_type(filename)
        if file_type is None:
            raise ValueError(f"Unsupported file type: {filename}")

        match file_type.extension:
            case ".pdf":
                return self.parse_pdf(file_bytes, on_page_progress=on_page_progress)
            case ".docx":
                return self._parse_docx_document(file_bytes, on_page_progress=on_page_progress)
            case ".txt":
                return self._parse_text_document(
                    file_bytes, markdown=False, on_page_progress=on_page_progress
                )
            case ".md":
                return self._parse_text_document(
                    file_bytes, markdown=True, on_page_progress=on_page_progress
                )
            case ".html" | ".htm":
                return self._parse_html_document(file_bytes, on_page_progress=on_page_progress)
            case ".png" | ".jpg" | ".jpeg" | ".webp":
                return self._parse_image_document(file_bytes, on_page_progress=on_page_progress)
            case _:
                raise ValueError(f"Unsupported file type: {filename}")
