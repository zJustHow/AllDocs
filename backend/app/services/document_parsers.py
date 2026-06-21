"""Parse non-PDF document formats into page rows."""

from __future__ import annotations

import html
import re
from io import BytesIO

from app.services.ingestion_types import PageRow
from app.services.pdf_page_text import is_toc_text

_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")


def decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Unable to decode text file")


def strip_html(raw_html: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "", raw_html)
    text = re.sub(r"(?s)<[^>]+>", "\n", text)
    text = html.unescape(text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def pages_with_bbox(
    pages: list[tuple[str | None, str, int]],
) -> list[PageRow]:
    return [(section, text, page, None, None) for section, text, page in pages]


def _flush_text_buffer(
    pages: list[tuple[str | None, str, int]],
    section: str | None,
    buffer: list[str],
    page_number: int = 1,
) -> None:
    if not buffer:
        return
    text = "\n".join(buffer).strip()
    if text and not is_toc_text(text):
        pages.append((section, text, page_number))


def parse_structured_text_pages(text: str, *, markdown: bool) -> list[tuple[str | None, str, int]]:
    text = text.strip()
    if not text:
        return []

    if not markdown:
        return [(None, text, 1)]

    pages: list[tuple[str | None, str, int]] = []
    section: str | None = None
    buffer: list[str] = []
    for line in text.splitlines():
        match = _MD_HEADING_RE.match(line)
        if match:
            _flush_text_buffer(pages, section, buffer)
            section = match.group(2).strip()
            buffer = []
            continue
        buffer.append(line)
    _flush_text_buffer(pages, section, buffer)
    return pages if pages else [(None, text, 1)]


def parse_docx_pages(file_bytes: bytes) -> list[tuple[str | None, str, int]]:
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(file_bytes))
    pages: list[tuple[str | None, str, int]] = []
    section: str | None = None
    buffer: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            _flush_text_buffer(pages, section, buffer)
            section = text
            buffer = []
            continue
        buffer.append(text)
    _flush_text_buffer(pages, section, buffer)

    if pages:
        return pages

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            pages.append((section, text, 1))
    return pages


def render_docx_preview_html(file_bytes: bytes) -> str:
    """Render DOCX text and tables as safe, self-contained preview HTML."""
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(file_bytes))
    content: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        escaped = html.escape(text)
        if style.startswith("heading"):
            match = re.search(r"(\d+)", style)
            level = min(6, max(1, int(match.group(1)))) if match else 2
            content.append(f"<h{level}>{escaped}</h{level}>")
        else:
            content.append(f"<p>{escaped}</p>")

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = "".join(f"<td>{html.escape(cell.text.strip())}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        if rows:
            content.append(f"<table>{''.join(rows)}</table>")

    body = "".join(content) or "<p></p>"
    return f"""<!doctype html><html><head><meta charset="utf-8"><style>
body{{margin:0;padding:32px;font:16px/1.65 system-ui,sans-serif;color:#222;background:#fff}}
h1,h2,h3,h4,h5,h6{{line-height:1.3;margin:1.2em 0 .55em}}p{{margin:.65em 0}}
table{{width:100%;border-collapse:collapse;margin:1em 0}}td{{border:1px solid #ccc;padding:8px;vertical-align:top}}
</style></head><body>{body}</body></html>"""
