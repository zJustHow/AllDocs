from io import BytesIO

from docx import Document

from app.services.document_parsers import render_docx_preview_html


def test_render_docx_preview_html_escapes_content_and_renders_tables() -> None:
    document = Document()
    document.add_heading("Safety <Guide>", level=1)
    document.add_paragraph("Disconnect power before service.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Step"
    table.cell(0, 1).text = "Action"
    buffer = BytesIO()
    document.save(buffer)

    preview = render_docx_preview_html(buffer.getvalue())

    assert "<h1>Safety &lt;Guide&gt;</h1>" in preview
    assert "<p>Disconnect power before service.</p>" in preview
    assert "<table>" in preview
    assert "<td>Action</td>" in preview
